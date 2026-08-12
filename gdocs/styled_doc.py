#!/usr/bin/env python3
"""styled_doc.py — a clean house style for native Google Docs.

Build a richly-styled .docx with python-docx (colored section headings, shaded
callout boxes, clean tables, embedded images, generous spacing), then upload it
to Google Drive AS a native Google Doc (Drive converts .docx -> Doc on upload),
OR PATCH an existing Doc in place so its share link / ID never changes.

WHY .docx and not markdown:
  Uploading markdown makes a *plain* Doc — it loses colors, callouts, and
  embedded images. The polished look — colored titles, lots of whitespace,
  digestible chunks, accent callouts — only survives the .docx -> Doc
  conversion. Use THIS module for styled briefs.

Credentials are PER-USER and never hard-coded here. Run the one-time setup first
(see SETUP.md):
  python3 scripts/auth.py
That stores your refresh token at ~/.config/google-skills/token.json and reads
your OAuth client from env vars GOOGLE_OAUTH_CLIENT_ID / GOOGLE_OAUTH_CLIENT_SECRET
or ~/.config/google-skills/oauth_client.json.

Typical use (see example_build.py for a full worked document):
  from styled_doc import Doc, upload_or_update
  d = Doc()
  d.title("Product Requirements Brief", "A HEALTH & SAFETY WRISTBAND", "Confidential — draft")
  d.heading("1", "Overview")
  d.body("One short, plain-English paragraph. Keep sentences tight.")
  d.bullet([("Lead-in: ", True, d.INK), ("the rest of the point.", False, d.BODY)])
  d.callout("Key point:", d.AMBER, "FEF3E2", "A highlighted box the reader can't miss.")
  d.image("/abs/path/to/pic.png", "A caption, italic and grey.", width=4.0)
  d.save("/abs/path/out.docx")
  print(upload_or_update("/abs/path/out.docx", title="My Brief"))          # new Doc
  print(upload_or_update("/abs/path/out.docx", doc_id="1aQD...Ke4BDJiSuU")) # update in place

SHARING IS AUTOMATIC: every upload/update also grants the workspace domain
(DEFAULT_SHARE_DOMAIN, "magicfp.com") writer-with-link access, so the URL this
returns can be pasted straight into chat. Pass share_domain=None for a private Doc.
"""
from __future__ import annotations

import json
import os
import sys
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- OAuth (Drive + Docs) — per-user, resolved at runtime, never hard-coded ----
CONFIG_DIR = Path.home() / ".config" / "google-skills"
TOKEN_PATH = CONFIG_DIR / "token.json"
CLIENT_PATH = CONFIG_DIR / "oauth_client.json"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GDOC_MIME = "application/vnd.google-apps.document"
SETUP_HINT = "Run the one-time setup:  python3 scripts/auth.py   (see SETUP.md)"

# Every Doc is shared with the workspace domain BY DEFAULT so the printed URL is
# immediately openable by the team ("Anyone in this group with the link can edit")
# and nobody has to click Share. Override with the env var, or pass
# share_domain=None to upload_or_update() for a genuinely private Doc.
DEFAULT_SHARE_DOMAIN = os.environ.get("GOOGLE_DOC_SHARE_DOMAIN", "magicfp.com")
_UNSET = object()


def _client_creds() -> tuple[str, str]:
    cid = os.environ.get("GOOGLE_OAUTH_CLIENT_ID")
    csec = os.environ.get("GOOGLE_OAUTH_CLIENT_SECRET")
    if cid and csec:
        return cid, csec
    if CLIENT_PATH.exists():
        data = json.loads(CLIENT_PATH.read_text())
        if "installed" in data:
            data = data["installed"]
        elif "web" in data:
            data = data["web"]
        if data.get("client_id") and data.get("client_secret"):
            return data["client_id"], data["client_secret"]
    sys.exit(
        "No Google OAuth client found.\n"
        f"Set GOOGLE_OAUTH_CLIENT_ID / GOOGLE_OAUTH_CLIENT_SECRET, or create {CLIENT_PATH}.\n"
        f"{SETUP_HINT}"
    )


def _access_token() -> str:
    if not TOKEN_PATH.exists():
        sys.exit(f"OAuth refresh token not found at {TOKEN_PATH}.\n{SETUP_HINT}")
    refresh = json.loads(TOKEN_PATH.read_text()).get("refresh_token")
    if not refresh:
        sys.exit(f"No refresh_token in {TOKEN_PATH}.\n{SETUP_HINT}")
    cid, csec = _client_creds()
    body = urllib.parse.urlencode({
        "client_id": cid,
        "client_secret": csec,
        "refresh_token": refresh,
        "grant_type": "refresh_token",
    }).encode()
    with urllib.request.urlopen(urllib.request.Request("https://oauth2.googleapis.com/token", data=body)) as r:
        return json.loads(r.read())["access_token"]


def upload_or_update(docx_path: str, *, title: str | None = None,
                     doc_id: str | None = None, folder_id: str | None = None,
                     share_domain=_UNSET, share_role: str = "writer") -> str:
    """Upload docx as a NEW native Doc (pass title), or convert-in-place over an
    existing Doc (pass doc_id). Returns the docs.google.com/document/d/<ID>/edit URL.

    Sharing: defaults to DEFAULT_SHARE_DOMAIN ("magicfp.com") on BOTH the create and
    the update path, so the returned URL is always ready to paste. Re-applying the
    same domain permission is idempotent (Drive returns the existing permission id).
    Pass share_domain=None to keep a Doc private."""
    token = _access_token()
    data = Path(docx_path).read_bytes()
    boundary = "===styled-doc-boundary==="
    meta: dict = {"mimeType": GDOC_MIME}
    if title:
        meta["name"] = title
    if folder_id and not doc_id:
        meta["parents"] = [folder_id]
    body = (
        f"--{boundary}\r\n"
        "Content-Type: application/json; charset=UTF-8\r\n\r\n"
        + json.dumps(meta) + "\r\n"
        f"--{boundary}\r\n"
        f"Content-Type: {DOCX_MIME}\r\n\r\n"
    ).encode() + data + f"\r\n--{boundary}--\r\n".encode()

    if doc_id:
        url = f"https://www.googleapis.com/upload/drive/v3/files/{doc_id}?uploadType=multipart&supportsAllDrives=true"
        method = "PATCH"
    else:
        url = "https://www.googleapis.com/upload/drive/v3/files?uploadType=multipart&supportsAllDrives=true"
        method = "POST"
    req = urllib.request.Request(url, data=body, method=method, headers={
        "Authorization": f"Bearer {token}",
        "Content-Type": f"multipart/related; boundary={boundary}",
    })
    with urllib.request.urlopen(req) as r:
        fid = json.loads(r.read())["id"]

    domain = DEFAULT_SHARE_DOMAIN if share_domain is _UNSET else share_domain
    if domain:
        share_with_domain(fid, domain, role=share_role, token=token)
    return f"https://docs.google.com/document/d/{fid}/edit"


def share_with_domain(file_id: str, domain: str = DEFAULT_SHARE_DOMAIN, *,
                      role: str = "writer", token: str | None = None) -> bool:
    """Give everyone in `domain` link-access to the file — the 'Anyone in this group
    with the link can edit' state. Idempotent: re-running returns the same permission.
    Never fatal — a sharing failure warns and leaves the (already uploaded) Doc intact."""
    perm = json.dumps({"type": "domain", "domain": domain, "role": role}).encode()
    req = urllib.request.Request(
        f"https://www.googleapis.com/drive/v3/files/{file_id}/permissions?supportsAllDrives=true",
        data=perm, method="POST",
        headers={"Authorization": f"Bearer {token or _access_token()}",
                 "Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req):
            return True
    except Exception as e:  # noqa: BLE001 — sharing must never lose the upload
        detail = e.read().decode()[:300] if hasattr(e, "read") else str(e)
        print(f"WARNING: could not share {file_id} with {domain}: {detail}\n"
              f"         The Doc exists but is PRIVATE — share it by hand before sending the link.",
              file=sys.stderr)
        return False


class Doc:
    """House-style document builder. Colors + fonts + spacing are baked in;
    you just call title/heading/body/bullet/numbered/callout/image/table/save."""

    # Palette — deep navy headings, red section numbers + accents, soft tints for callouts.
    BLUE_DK = RGBColor(0x1E, 0x3A, 0x8A)
    BLUE    = RGBColor(0x1D, 0x4E, 0xD8)
    RED     = RGBColor(0xDC, 0x26, 0x26)
    GREEN   = RGBColor(0x05, 0x96, 0x69)
    AMBER   = RGBColor(0xB4, 0x53, 0x09)
    INK     = RGBColor(0x11, 0x18, 0x27)   # near-black, for bold lead-ins
    BODY    = RGBColor(0x37, 0x41, 0x51)   # default body grey-ink
    GRAY    = RGBColor(0x6B, 0x72, 0x80)   # captions / notes
    WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
    FONT = "Arial"

    def __init__(self):
        self.doc = Document()
        for s in self.doc.sections:
            s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
            s.left_margin = Inches(0.85); s.right_margin = Inches(0.85)
        n = self.doc.styles["Normal"]
        n.font.name = self.FONT
        n.font.size = Pt(10.5)
        n.font.color.rgb = self.BODY
        n.paragraph_format.line_spacing = 1.22     # airy, not cramped
        n.paragraph_format.space_after = Pt(6)

    # ---- low-level run ----
    def run(self, p, text, size=10.5, color=None, bold=False, italic=False):
        r = p.add_run(text)
        r.font.name = self.FONT; r.font.size = Pt(size)
        r.font.bold = bold; r.font.italic = italic
        r.font.color.rgb = color if color is not None else self.BODY
        return r

    # ---- header block (document title + red subtitle + grey tagline + rule) ----
    def title(self, title, subtitle=None, tagline=None):
        t = self.doc.add_paragraph(); t.paragraph_format.space_after = Pt(2)
        self.run(t, title, size=25, color=self.BLUE_DK, bold=True)
        if subtitle:
            s = self.doc.add_paragraph(); s.paragraph_format.space_after = Pt(4)
            self.run(s, subtitle.upper(), size=11.5, color=self.RED, bold=True)
        tg = self.doc.add_paragraph(); tg.paragraph_format.space_after = Pt(10)
        if tagline:
            self.run(tg, tagline, size=9.5, color=self.GRAY, italic=True)
        self._bottom_border(tg, "1E3A8A", 14)
        return t

    # ---- numbered section heading with thin underline ----
    def heading(self, num, text):
        p = self.doc.add_paragraph()
        # generous air ABOVE each section — let sections breathe
        p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(6)
        self.run(p, f"{num}   ", size=14, color=self.RED, bold=True)
        self.run(p, text, size=14, color=self.BLUE_DK, bold=True)
        self._bottom_border(p, "DBEAFE", 8)
        return p

    def body(self, text, size=10.5, color=None, space_after=6):
        """A single-colour paragraph. If the paragraph needs an emphasised lead-in,
        a highlighted phrase, or a link, use rich() instead — don't fight body()."""
        p = self.doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
        self.run(p, text, size=size, color=color if color is not None else self.BODY)
        return p

    def rich(self, parts, space_after=8):
        """A body paragraph built from mixed runs — the workhorse for colour + links.
        Each seg is either:
          (text, bold, color)            -> a coloured/bold run, e.g. ("first words", True, self.RED)
          ("LINK", text, url[, hex])     -> an inline hyperlink (default blue, underlined)
        Use this for the opening of a paragraph in colour, to highlight the key phrase
        mid-sentence, and to link any product/competitor name."""
        p = self.doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
        self._segs(p, parts)
        return p

    def note(self, lead, text):
        """Small grey footnote-style line: bold blue lead-in + grey body."""
        p = self.doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        self.run(p, lead + "  ", size=9.5, color=self.BLUE_DK, bold=True)
        self.run(p, text, size=9.5, color=self.GRAY)
        return p

    def bullet(self, parts):
        """parts = list of (text, bold, color) tuples — or ('LINK', text, url) — on one line.
        Give every bullet a COLOURED bold lead-in (vary RED/BLUE_DK/GREEN/AMBER by theme),
        then the detail in BODY grey."""
        p = self.doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(4)
        self._segs(p, parts)
        return p

    def numbered(self, parts):
        p = self.doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(5)
        self._segs(p, parts)
        return p

    def callout(self, lead, lead_color, fill_hex, text):
        """Shaded one-cell box. lead is bold colored; text is INK. fill_hex like 'FEF3E2'."""
        t = self.doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._no_table_borders(t)
        cell = t.cell(0, 0); self._shade(cell, fill_hex); self._cell_margins(cell, 120, 120, 160, 160)
        cp = cell.paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
        if lead:
            self.run(cp, lead + "  ", size=10.5, color=lead_color, bold=True)
        self.run(cp, text, size=10.5, color=self.INK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def image(self, path, caption=None, width=6.0):
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(path, width=Inches(width))
        if caption:
            c = self.doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_after = Pt(8)
            self.run(c, caption, size=8.5, color=self.GRAY, italic=True)

    def table(self, headers, rows):
        """Navy header row + zebra body. headers=[..]; rows=[[c0,c1,..],..].
        First column is rendered bold INK, others BODY."""
        tbl = self.doc.add_table(rows=1, cols=len(headers)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        for i, label in enumerate(headers):
            self._shade(hdr[i], "1E3A8A"); self._cell_margins(hdr[i])
            hp = hdr[i].paragraphs[0]; hp.paragraph_format.space_after = Pt(0)
            self.run(hp, label, size=10, color=self.WHITE, bold=True)
        for ri, row in enumerate(rows):
            cells = tbl.add_row().cells
            fill = "F3F6FB" if ri % 2 == 0 else "FFFFFF"
            for ci, val in enumerate(row):
                self._shade(cells[ci], fill); self._cell_margins(cells[ci])
                cp = cells[ci].paragraphs[0]; cp.paragraph_format.space_after = Pt(0)
                self.run(cp, val, size=9.5, color=(self.INK if ci == 0 else self.BODY), bold=(ci == 0))
        b = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "D9DEE7"); b.append(e)
        tbl._tbl.tblPr.append(b)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return tbl

    def footer(self, text):
        f = self.doc.add_paragraph(); f.paragraph_format.space_before = Pt(10)
        self.run(f, text, size=8.5, color=self.GRAY, italic=True)

    def save(self, path):
        self.doc.save(path)
        return path

    # ---- run-segment emitter (shared by rich/bullet/numbered) ----
    def _segs(self, p, parts):
        for seg in parts:
            if seg and seg[0] == "LINK":               # ("LINK", text, url[, hex])
                text, url = seg[1], seg[2]
                color = seg[3] if len(seg) > 3 else "1D4ED8"
                self._hyperlink(p, text, url, color)
            else:                                       # (text, bold, color)
                text, bold, color = seg
                self.run(p, text, size=10.5, color=color, bold=bold)

    def _hyperlink(self, p, text, url, color="1D4ED8"):
        """Real clickable hyperlink run (underlined, blue) appended to paragraph p."""
        rid = p.part.relate_to(
            url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True)
        h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), rid)
        r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), self.FONT); rf.set(qn("w:hAnsi"), self.FONT); rPr.append(rf)
        col = OxmlElement("w:color"); col.set(qn("w:val"), color); rPr.append(col)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "21"); rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; r.append(t)
        h.append(r); p._p.append(h)
        return h

    # ---- XML helpers (cell shading, margins, borders) ----
    def _shade(self, cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc)
        tcPr.append(shd)

    def _cell_margins(self, cell, top=80, bottom=80, left=120, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        m = OxmlElement("w:tcMar")
        for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
            e = OxmlElement(f"w:{tag}"); e.set(qn("w:w"), str(val)); e.set(qn("w:type"), "dxa"); m.append(e)
        tcPr.append(m)

    def _no_table_borders(self, table):
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"), "none"); borders.append(e)
        table._tbl.tblPr.append(borders)

    def _bottom_border(self, paragraph, hexc="1E3A8A", size=10):
        pPr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "6"); b.set(qn("w:color"), hexc)
        pbdr.append(b); pPr.append(pbdr)
