#!/usr/bin/env python3
"""Deterministic markdown -> house-styled docx converter for the planning docs.

Maps our docs' markdown onto styled_doc.Doc primitives: numbered navy/red headings,
zebra tables, monospace code boxes (ASCII diagrams survive), inline bold/code/links,
blockquote callouts, and a compact small-type Sources section.
"""
from __future__ import annotations

import re
import sys

from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from styled_doc import Doc
from md_to_html import preprocess

MONO = "Consolas"
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))")
LINK = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
URL = re.compile(r"(https?://\S+)")
HNUM = re.compile(r"^(\d+(?:\.\d+)*)[.)]?\s+(.*)$")


def mono_run(d, p, text, size=8.0, color=None):
    r = p.add_run(text)
    r.font.name = MONO
    r._element.rPr.rFonts.set(qn("w:cs"), MONO)
    r.font.size = Pt(size)
    r.font.color.rgb = color if color is not None else d.INK
    return r


def emit_inline(d, p, text, size=10.5, first_bold_color=None, base_color=None):
    """Emit mixed bold/code/link segments into paragraph p."""
    first = True
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            color = first_bold_color if (first and first_bold_color is not None) else d.INK
            d.run(p, tok[2:-2], size=size, color=color, bold=True)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2 and not tok.startswith("**"):
            d.run(p, tok[1:-1], size=size, color=base_color if base_color is not None else d.BODY, italic=True)
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            mono_run(d, p, tok[1:-1], size=size - 1.0)
        else:
            m = LINK.fullmatch(tok)
            if m:
                d._hyperlink(p, m.group(1), m.group(2))
            else:
                d.run(p, tok, size=size, color=base_color if base_color is not None else d.BODY)
        if tok.strip():
            first = False


def plain(text):
    """Flatten inline markdown for table cells."""
    text = LINK.sub(lambda m: m.group(1), text)
    return text.replace("**", "").replace("`", "").strip()


def code_block(d, lines):
    t = d.doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4"); e.set(qn("w:color"), "D9DEE7")
        b.append(e)
    t._tbl.tblPr.append(b)
    cell = t.cell(0, 0)
    d._shade(cell, "F3F4F6")
    d._cell_margins(cell, 100, 100, 140, 140)
    first = True
    for ln in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        mono_run(d, p, ln if ln else " ")
    d.doc.add_paragraph().paragraph_format.space_after = Pt(2)


def subheading(d, text, level):
    p = d.doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 3 else 10)
    p.paragraph_format.space_after = Pt(4)
    m = HNUM.match(text)
    if level == 3 and m:
        d.run(p, m.group(1) + "  ", size=12, color=d.RED, bold=True)
        d.run(p, plain(m.group(2)), size=12, color=d.BLUE_DK, bold=True)
    else:
        d.run(p, plain(text), size=11.5 if level == 3 else 11, color=d.BLUE_DK if level == 3 else d.INK, bold=True)


def convert(md_path, out_path):
    raw = preprocess(open(md_path, encoding="utf-8").read().splitlines())
    d = Doc()
    i, n = 0, len(raw)
    seen_title = False
    section = 0
    in_sources = False

    def peek_subtitle(j):
        # bold standalone "Document 0X of 07 ..." line following the H1
        while j < n and not raw[j].strip():
            j += 1
        if j < n and raw[j].strip().startswith("**") and raw[j].strip().endswith("**"):
            return raw[j].strip().strip("*").strip(), j + 1
        return None, j

    while i < n:
        line = raw[i]
        s = line.strip()

        if s.startswith("```"):
            i += 1
            block = []
            while i < n and not raw[i].strip().startswith("```"):
                block.append(raw[i])
                i += 1
            code_block(d, block)
            i += 1
            continue

        if s.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?$", raw[i + 1].strip()):
            headers = [plain(c) for c in s.strip("|").split("|")]
            i += 2
            rows = []
            while i < n and raw[i].strip().startswith("|"):
                rows.append([plain(c) for c in raw[i].strip().strip("|").split("|")])
                i += 1
            width = len(headers)
            rows = [(r + [""] * width)[:width] for r in rows]
            d.table(headers, rows)
            continue

        if s.startswith("#### "):
            subheading(d, s[5:], 4)
        elif s.startswith("### "):
            subheading(d, s[4:], 3)
        elif s.startswith("## "):
            text = plain(s[3:])
            m = HNUM.match(text)
            in_sources = text.lower().startswith("sources")
            if m:
                section = int(m.group(1).split(".")[0])
                d.heading(m.group(1), m.group(2))
            else:
                section += 1
                d.heading(str(section), text)
        elif s.startswith("# ") and not seen_title:
            seen_title = True
            sub, j = peek_subtitle(i + 1)
            tagline = None
            while j < n and not raw[j].strip():
                j += 1
            if j < n and not raw[j].strip().startswith(("#", "|", "-", "*", ">", "```")) and raw[j].strip() != "---":
                tag_lines = []
                while j < n and raw[j].strip() and not raw[j].strip().startswith(("#", "|", ">", "```")) and raw[j].strip() != "---":
                    tag_lines.append(plain(raw[j].strip()))
                    j += 1
                tagline = " ".join(tag_lines)
            d.title(plain(s[2:]), sub, tagline)
            i = j
            continue
        elif s == "---" or s == "":
            pass
        elif s.startswith(">"):
            d.callout("", d.BLUE_DK, "DBEAFE", plain(s.lstrip("> ")))
        elif re.match(r"^\s*[-*]\s+", line):
            indent = len(line) - len(line.lstrip())
            content = re.sub(r"^\s*[-*]\s+", "", line)
            p = d.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            if indent >= 2:
                d.run(p, "–  ", size=10.5, color=d.GRAY)
            emit_inline(d, p, content, first_bold_color=d.BLUE_DK)
        elif re.match(r"^\s*\d+\.\s+", line):
            content = re.sub(r"^\s*\d+\.\s+", "", line)
            if in_sources:
                num = re.match(r"^\s*(\d+)\.", line).group(1)
                p = d.doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                d.run(p, num + ". ", size=8.5, color=d.GRAY, bold=True)
                um = URL.search(content)
                if um:
                    pre = content[:um.start()].strip()
                    post = content[um.end():].strip()
                    if pre:
                        d.run(p, plain(pre) + " ", size=8.5, color=d.GRAY)
                    d._hyperlink(p, um.group(1).rstrip(".,)"), um.group(1).rstrip(".,)"))
                    if post:
                        d.run(p, " " + plain(post), size=8.5, color=d.GRAY)
                else:
                    d.run(p, plain(content), size=8.5, color=d.GRAY)
            else:
                p = d.doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(4)
                emit_inline(d, p, content, first_bold_color=d.BLUE_DK)
        else:
            p = d.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            emit_inline(d, p, s, first_bold_color=d.BLUE_DK)
        i += 1

    d.save(out_path)
    return out_path


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
    print("wrote", sys.argv[2])
